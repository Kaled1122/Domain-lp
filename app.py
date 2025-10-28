import os, io, re, logging
from datetime import datetime
from flask import Flask, request, jsonify, send_file
from flask_cors import CORS
from psycopg2.pool import SimpleConnectionPool
from PyPDF2 import PdfReader
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
from reportlab.lib.styles import getSampleStyleSheet
from reportlab.lib.pagesizes import A4
from openai import OpenAI
from io import BytesIO
from docx import Document
from docx.shared import Inches
from docx.enum.section import WD_ORIENT

# ------------------------------------------------------------
# APP SETUP
# ------------------------------------------------------------
app = Flask(__name__)
CORS(app, resources={r"/*": {"origins": "*"}})
logging.basicConfig(level=logging.INFO)

# Load environment variables
OPENAI_KEY = os.getenv("OPENAI_API_KEY")
DB_URL = os.getenv("DATABASE_URL")

if not OPENAI_KEY:
    logging.error("❌ Missing OPENAI_API_KEY — please set it in Railway environment variables.")
if not DB_URL:
    logging.warning("⚠️ No DATABASE_URL found — DB features disabled.")

# Initialize OpenAI client
client = OpenAI(api_key=OPENAI_KEY)

# Optional DB connection pool
pool = SimpleConnectionPool(1, 10, dsn=DB_URL) if DB_URL else None

# ------------------------------------------------------------
# HELPERS
# ------------------------------------------------------------
def get_conn():
    """Get connection from pool (if available)."""
    if not pool:
        raise Exception("Database not configured.")
    return pool.getconn()

def put_conn(c):
    """Return connection to pool."""
    if pool and c:
        pool.putconn(c)

def extract_text_from_pdf(file):
    """Extracts up to 5000 characters of text from uploaded PDF."""
    try:
        reader = PdfReader(file)
        text = "\n".join([p.extract_text() or "" for p in reader.pages])
        return text[:5000] if text.strip() else "PDF content extracted (empty)."
    except Exception as e:
        logging.error(f"PDF extraction failed: {e}")
        return "PDF uploaded (text extraction failed)."

def safe_float(v):
    """Convert to float safely (returns 0.0 if blank or invalid)."""
    try:
        if v in [None, "", " "]:
            return 0.0
        return float(v)
    except Exception:
        return 0.0

# ------------------------------------------------------------
# LESSON PLAN GENERATOR
# ------------------------------------------------------------
def generate_lesson_plan_text(teacher, title, duration, cefr, profile, content):
    """Generate structured HTML lesson plan including domain checklists."""
    prompt = f"""
You are a senior English Language Teaching (ELT) instructional designer for BAE Systems.

Generate a professional, structured HTML lesson plan for classroom delivery using a formal instructional tone.

Requirements:
- Output valid HTML (no markdown, no code fences).
- Include ALL sections as shown below.

<h2>Lesson Plan</h2>
<b>Title:</b> {title}<br>
<b>Teacher:</b> {teacher}<br>
<b>Duration:</b> {duration}<br>
<b>CEFR Level:</b> {cefr}<br>
<b>Learner Profile:</b> {profile}<br>

<h3>1. Lesson Objectives</h3>
<ul>
<li>At least 3 measurable objectives aligned with CEFR outcomes.</li>
</ul>

<h3>2. Lesson Plan Structure</h3>
<table cellspacing="0" cellpadding="6" style="border-collapse:collapse;width:100%">
<tr style="background:#e0e7ff;font-weight:bold;text-align:left">
<th>Stage</th><th>Duration</th><th>Objective/Skill</th><th>Activities</th><th>Teacher Role</th><th>Learner Role</th><th>Materials</th>
</tr>
<tr><td>Warm-Up</td><td>5–10 min</td><td>Activate prior knowledge.</td><td>Quick discussion or brainstorming.</td><td>Facilitates, elicits responses.</td><td>Responds, shares ideas.</td><td>Board, visuals.</td></tr>
<tr><td>Presentation</td><td>10 min</td><td>Introduce target language.</td><td>Demonstrate grammar or vocabulary use in context.</td><td>Explains, models, checks understanding.</td><td>Listens and takes notes.</td><td>Projector, whiteboard.</td></tr>
<tr><td>Practice (Controlled)</td><td>10 min</td><td>Reinforce comprehension.</td><td>Pair or group structured exercises.</td><td>Monitors, corrects errors.</td><td>Completes tasks accurately.</td><td>Worksheets.</td></tr>
<tr><td>Production (Freer)</td><td>10 min</td><td>Apply language creatively.</td><td>Role-play or group presentation.</td><td>Guides, observes, supports.</td><td>Uses target language independently.</td><td>Real-life prompts.</td></tr>
<tr><td>Review & Wrap-Up</td><td>5 min</td><td>Summarize learning outcomes.</td><td>Class recap and reflection.</td><td>Summarizes key points.</td><td>Reflects, asks questions.</td><td>Notebook.</td></tr>
</table>

<h3>3. Supporting Details</h3>
<table cellspacing="0" cellpadding="6" style="border-collapse:collapse;width:100%">
<tr style="background:#e0e7ff;font-weight:bold;text-align:left">
<th>Stage</th>
<th>Purpose</th>
<th>Method</th>
<th>Expected Outcome</th>
</tr>
<tr><td>Warm-Up</td><td>...</td><td>...</td><td>...</td></tr>
<tr><td>Presentation</td><td>...</td><td>...</td><td>...</td></tr>
<tr><td>Practice (Controlled)</td><td>...</td><td>...</td><td>...</td></tr>
<tr><td>Production (Freer)</td><td>...</td><td>...</td><td>...</td></tr>
<tr><td>Review & Wrap-Up</td><td>...</td><td>...</td><td>...</td></tr>
</table>

<h3>4. Performance Domain Checklists</h3>
<div style="background:#f9fafb;border-left:4px solid #2563eb;padding:10px;margin:10px 0;">
  <b>Each domain checklist</b> is rated <b>1–5</b> (5 = Excellent, 1 = Poor). Each domain totals <b>25 points</b>.
</div>

Lesson content reference:
{content}

<h4 style="color:#2563eb">Understanding (U)</h4>
<ul><li>...</li><li>...</li></ul>

<h4 style="color:#16a34a">Application (A)</h4>
<ul><li>...</li><li>...</li></ul>

<h4 style="color:#f59e0b">Communication (C)</h4>
<ul><li>...</li><li>...</li></ul>

<h4 style="color:#dc2626">Behavior (B)</h4>
<ul><li>...</li><li>...</li></ul>

<h3>5. Score Interpretation Key</h3>
<table cellspacing="0" cellpadding="6" style="border-collapse:collapse;width:50%">
<tr style="background:#059669;color:white;font-weight:bold;text-align:center"><th>Score Range</th><th>Performance Level</th></tr>
<tr><td style="text-align:center">90–100</td><td>Outstanding</td></tr>
<tr><td style="text-align:center">75–89</td><td>Good</td></tr>
<tr><td style="text-align:center">50–74</td><td>Developing</td></tr>
<tr><td style="text-align:center">0–49</td><td>Needs Improvement</td></tr>
</table>

<h3>6. Reflection (Instructor Review)</h3>
<ul><!-- The model will generate three context-specific reflection questions below --></ul>
"""
    try:
        response = client.chat.completions.create(
            model="gpt-4o-mini",
            temperature=0.4,
            messages=[
                {
                    "role": "system",
                    "content": (
                        "You are an expert ELT instructional designer working under BAE Systems standards. "
                        "Do not restate instructions or placeholders. "
                        "Only output fully populated HTML sections that match the requested structure."
                        "In the Performance Domain Checklists section, generate 5 measurable, lesson-specific indicators "
                        "for each domain (Understanding, Application, Communication, Behavior), "
                        "based strictly on the uploaded lesson content."
                    ),
                },
                {
                    "role": "user",
                    "content": f"Generate the following HTML structure exactly, filling all fields with relevant content derived from the uploaded lesson:\n\n{prompt}"
                },
            ],
        )

        html = response.choices[0].message.content.strip()
        html = re.sub(r"^```(?:html)?|```$", "", html, flags=re.MULTILINE).strip()
        html = html.replace("\\n", "\n").replace('\\"', '"')
        return html

    except Exception as e:
        logging.error(f"AI generation failed: {e}")
        return f"<p style='color:red'>AI generation failed: {e}</p>"

# ------------------------------------------------------------
# NEW: WORD (DOCX) DOWNLOAD ENDPOINT
# ------------------------------------------------------------
# ------------------------------------------------------------
# NEW: WORD (DOCX) DOWNLOAD ENDPOINT (IMPROVED FORMATTED VERSION)
# ------------------------------------------------------------
@app.post("/download_lesson_docx")
def download_lesson_docx():
    """Converts generated HTML to a formatted Word document (landscape)."""
    try:
        html_content = request.form.get("html", "")
        if not html_content:
            return jsonify({"status": "error", "message": "No HTML received."}), 400

        # ✅ Create a Word doc and set landscape layout
        doc = Document()
        section = doc.sections[0]
        section.orientation = WD_ORIENT.LANDSCAPE
        section.page_width, section.page_height = section.page_height, section.page_width

        # ✅ Convert basic HTML formatting
        html_content = html_content.replace("<br>", "\n").replace("</p>", "\n\n")

        # Split by major blocks (headings, tables, lists, etc.)
        html_blocks = re.split(r'(<h\d[^>]*>.*?</h\d>|<table.*?</table>|<ul>.*?</ul>)', html_content, flags=re.S | re.I)

        for block in html_blocks:
            if not block.strip():
                continue

            # ===== HEADINGS =====
            if re.match(r"<h2", block, re.I):
                doc.add_heading(re.sub(r"<.*?>", "", block), level=1)
            elif re.match(r"<h3", block, re.I):
                h = doc.add_heading(re.sub(r"<.*?>", "", block), level=2)
                for r in h.runs:
                    r.bold = True
            elif re.match(r"<h4", block, re.I):
                h = doc.add_heading(re.sub(r"<.*?>", "", block), level=3)
                for r in h.runs:
                    r.bold = True

            # ===== TABLES =====
            elif "<table" in block.lower():
                rows = re.findall(r"<tr.*?>(.*?)</tr>", block, flags=re.S)
                if rows:
                    first_row = re.findall(r"<t[hd].*?>(.*?)</t[hd]>", rows[0], flags=re.S)
                    table = doc.add_table(rows=1, cols=len(first_row))
                    hdr_cells = table.rows[0].cells
                    for i, cell in enumerate(first_row):
                        text = re.sub(r"<.*?>", "", cell).strip()
                        hdr_cells[i].text = text
                        for p in hdr_cells[i].paragraphs:
                            for r in p.runs:
                                r.bold = True
                    for row in rows[1:]:
                        cols = re.findall(r"<t[hd].*?>(.*?)</t[hd]>", row, flags=re.S)
                        if not cols:
                            continue
                        cells = table.add_row().cells
                        for i, cell in enumerate(cols):
                            text = re.sub(r"<.*?>", "", cell).strip()
                            cells[i].text = text

            # ===== BULLET LISTS =====
            elif "<ul" in block.lower():
                items = re.findall(r"<li.*?>(.*?)</li>", block, flags=re.S)
                for li in items:
                    doc.add_paragraph(re.sub(r"<.*?>", "", li), style="List Bullet")

            # ===== PARAGRAPHS =====
            else:
                text = re.sub(r"<.*?>", "", block).strip()
                if text:
                    p = doc.add_paragraph(text)
                    for r in p.runs:
                        r.font.name = "Arial"

        # ✅ Save final document
        output = BytesIO()
        doc.save(output)
        output.seek(0)

        return send_file(
            output,
            mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            as_attachment=True,
            download_name="lesson_plan.docx"
        )

    except Exception as e:
        logging.error(f"❌ DOCX generation failed: {e}")
        return jsonify({"status": "error", "message": str(e)}), 500

# ------------------------------------------------------------
# ROUTES (unchanged)
# ------------------------------------------------------------
@app.route("/", methods=["GET"])
def home():
    return jsonify({"status": "OK", "message": "Lesson Planner Backend Active"})


@app.post("/generate_lesson")
def generate_lesson():
    """Handles lesson plan generation requests from the frontend."""
    try:
        f = request.form
        teacher = f.get("teacher", "")
        title = f.get("lesson_title", "")
        duration = f.get("duration", "")
        cefr = f.get("cefr", "")
        profile = f.get("profile", "")
        file = request.files.get("file")

        content = ""
        if file:
            if file.filename.lower().endswith(".pdf"):
                content = extract_text_from_pdf(file)
            else:
                content = file.read().decode("utf-8", errors="ignore")

        html_output = generate_lesson_plan_text(teacher, title, duration, cefr, profile, content)
        return jsonify({"status": "success", "html": html_output})

    except Exception as e:
        logging.error(f"❌ Error in /generate_lesson: {e}")
        return jsonify({"status": "error", "message": str(e)}), 500


if __name__ == "__main__":
    PORT = int(os.environ.get("PORT", 5000))
    app.run(host="0.0.0.0", port=PORT)
